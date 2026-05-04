"""
Generate real files (docx, xlsx, pdf, zip, code, text) from AI response markers.

Format the AI uses:
    <<<FILE:nome.ext>>>
    conteúdo do arquivo
    <<<END_FILE>>>
"""
import csv
import io
import re
import zipfile
from pathlib import Path
from typing import Any
from xml.sax.saxutils import escape as xml_escape

FILE_BLOCK_RE = re.compile(r'<<<FILE:([^\n>]+)>>>\n?(.*?)<<<END_FILE>>>', re.DOTALL)
DISPLAY_FILE_RE = re.compile(
    r'^\[📄\s*([^\]\n]+?\.[A-Za-z0-9]{2,8})\]\s*\n(.*?)(?=^\[📄\s*[^\]\n]+?\.[A-Za-z0-9]{2,8}\]\s*$|^\[📦|\n\*\*Para (?:usar|converter)|\nPara (?:usar|converter)|\nQuer que|\nOs arquivos|\Z)',
    re.DOTALL | re.MULTILINE,
)

# File-creation hint injected into the system prompt
FILE_CREATION_HINT = (
    "\n\nCapacidade de criar arquivos: quando solicitado, use EXATAMENTE:\n"
    "<<<FILE:nome.extensao>>>\n"
    "conteúdo completo aqui\n"
    "<<<END_FILE>>>\n"
    "Não escreva '[📄 arquivo]' como texto, não diga para converter manualmente e não entregue apenas markdown.\n"
    "O arquivo real será criado automaticamente a partir do bloco <<<FILE>>>.\n"
    "Extensões: .txt .py .js .ts .html .css .json .md .csv .xml .yaml .sh .sql "
    ".docx .xlsx .pptx .pdf e qualquer linguagem de programação.\n"
    "Para .docx/.pdf use markdown (# Título, ## Seção, **negrito**, - lista).\n"
    "Para .xlsx use CSV (col1,col2\\nval1,val2).\n"
    "Para .pptx use markdown: separe slides com ---; use # para título e - para bullets. "
    "Crie títulos curtos, bullets objetivos e divida o conteúdo em slides limpos; "
    "o sistema aplicará automaticamente layouts profissionais, cards, métricas, fluxos e elementos visuais genéricos.\n"
    "Pode criar múltiplos arquivos em uma resposta (serão compactados em ZIP automaticamente)."
)

FILE_KEYWORDS = (
    "cri", "gera", "escreve", "produz", "arquivo", "pdf", "docx", "excel",
    "planilha", "script", "código", "codigo", "zip", "relatório", "relatorio",
    "documento", "apresentação", "apresentacao", "slides", "powerpoint",
    "power point", "pptx",
    "file", "create", "generate",
)


def wants_file_creation(prompt: str) -> bool:
    low = prompt.lower()
    return any(kw in low for kw in FILE_KEYWORDS)


# ---------------------------------------------------------------------------
# Parsing
# ---------------------------------------------------------------------------

def parse_file_blocks(text: str) -> tuple[str, list[dict[str, str]]]:
    """Extract <<<FILE:name>>>...<<<END_FILE>>> blocks.
    Returns (cleaned_text, [{filename, content}, ...]).
    """
    files: list[dict[str, str]] = []

    def _replace(m: re.Match) -> str:
        filename = m.group(1).strip()
        content = m.group(2).rstrip("\n")
        files.append({"filename": filename, "content": content})
        return f"[📄 {filename}]"

    cleaned = FILE_BLOCK_RE.sub(_replace, text)
    if files:
        return _cleanup_file_answer(cleaned), files

    def _replace_display(m: re.Match) -> str:
        filename = m.group(1).strip()
        content = _trim_display_file_content(m.group(2))
        if content:
            files.append({"filename": filename, "content": content})
        return f"[📄 {filename}]"

    cleaned = DISPLAY_FILE_RE.sub(_replace_display, text)
    if files:
        cleaned = _cleanup_file_answer(cleaned)
    return cleaned, files


def _trim_display_file_content(content: str) -> str:
    lines = content.strip().splitlines()
    cut_at = len(lines)
    stop_patterns = (
        r"^\*\*Para (?:usar|converter)",
        r"^Para (?:usar|converter)",
        r"^Quer que\b",
        r"^Os arquivos\b",
        r"^Todos os arquivos\b",
        r"^Instruções:",
    )
    for i, line in enumerate(lines):
        if any(re.match(pattern, line.strip(), re.IGNORECASE) for pattern in stop_patterns):
            cut_at = i
            break
    return "\n".join(lines[:cut_at]).strip()


def _cleanup_file_answer(text: str) -> str:
    lines = text.strip().splitlines()
    cleaned: list[str] = []
    stop_patterns = (
        r"^\*\*Para (?:usar|converter)",
        r"^Para (?:usar|converter)",
        r"^Instruções:",
        r"^Quer que eu ajuste",
        r"^Quer que eu crie",
    )
    for line in lines:
        if any(re.match(pattern, line.strip(), re.IGNORECASE) for pattern in stop_patterns):
            break
        cleaned.append(line)
    return "\n".join(cleaned).strip()


# ---------------------------------------------------------------------------
# File generators
# ---------------------------------------------------------------------------

def _safe_name(filename: str) -> str:
    return re.sub(r"[^\w.\-]", "_", filename)


def _create_text(content: str, path: Path) -> None:
    path.write_text(content, encoding="utf-8")


def _create_docx(content: str, path: Path) -> None:
    from docx import Document

    doc = Document()
    for line in content.split("\n"):
        s = line.rstrip()
        if s.startswith("### "):
            doc.add_heading(s[4:], level=3)
        elif s.startswith("## "):
            doc.add_heading(s[3:], level=2)
        elif s.startswith("# "):
            doc.add_heading(s[2:], level=1)
        elif s.startswith(("- ", "* ")):
            doc.add_paragraph(s[2:], style="List Bullet")
        elif re.match(r"^\d+\. ", s):
            doc.add_paragraph(re.sub(r"^\d+\. ", "", s), style="List Number")
        else:
            p = doc.add_paragraph()
            for part in re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*)", s):
                if part.startswith("**") and part.endswith("**"):
                    p.add_run(part[2:-2]).bold = True
                elif part.startswith("*") and part.endswith("*"):
                    p.add_run(part[1:-1]).italic = True
                else:
                    p.add_run(part)
    doc.save(str(path))


def _create_xlsx(content: str, path: Path) -> None:
    import openpyxl

    wb = openpyxl.Workbook()
    ws = wb.active
    reader = csv.reader(io.StringIO(content))
    for row in reader:
        ws.append(row)
    wb.save(str(path))


def _create_pdf(content: str, path: Path) -> None:
    from fpdf import FPDF
    from fpdf.enums import XPos, YPos

    pdf = FPDF()
    pdf.set_margins(20, 20, 20)
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_font("Helvetica", "", 11)
    # Use fixed width — multi_cell(w=0) in fpdf2 2.8 means "remaining width",
    # which becomes 0 after the first call. Must use explicit printable width.
    W = pdf.w - pdf.l_margin - pdf.r_margin
    kw = {"new_x": XPos.LMARGIN, "new_y": YPos.NEXT}

    for line in content.split("\n"):
        s = line.rstrip()
        plain = re.sub(r"\*{1,2}([^*]+)\*{1,2}", r"\1", s)
        safe = plain.encode("latin-1", errors="replace").decode("latin-1")

        if s.startswith("# "):
            pdf.set_font("Helvetica", "B", 16)
            pdf.multi_cell(W, 8, safe[2:], **kw)
            pdf.set_font("Helvetica", "", 11)
        elif s.startswith("## "):
            pdf.set_font("Helvetica", "B", 13)
            pdf.multi_cell(W, 7, safe[3:], **kw)
            pdf.set_font("Helvetica", "", 11)
        elif s.startswith("### "):
            pdf.set_font("Helvetica", "B", 11)
            pdf.multi_cell(W, 6, safe[4:], **kw)
            pdf.set_font("Helvetica", "", 11)
        elif not s:
            pdf.ln(4)
        else:
            pdf.multi_cell(W, 6, safe, **kw)

    pdf.output(str(path))


def _clean_markdown_text(text: str) -> str:
    text = re.sub(r"^\s{0,3}#{1,6}\s+", "", text.strip())
    text = re.sub(r"^\s*(?:[-*]|\d+\.)\s+", "", text)
    text = re.sub(r"\*{1,2}([^*]+)\*{1,2}", r"\1", text)
    return text.strip()


def _split_pptx_slides(content: str) -> list[dict[str, list[str] | str]]:
    lines = content.replace("\r\n", "\n").split("\n")
    parts: list[list[str]] = []

    if any(re.match(r"^\s*---+\s*$", line) for line in lines):
        current: list[str] = []
        for line in lines:
            if re.match(r"^\s*---+\s*$", line):
                if any(x.strip() for x in current):
                    parts.append(current)
                current = []
            else:
                current.append(line)
        if any(x.strip() for x in current):
            parts.append(current)
    else:
        current = []
        for line in lines:
            if re.match(r"^\s*#\s+", line) and any(x.strip() for x in current):
                parts.append(current)
                current = [line]
            else:
                current.append(line)
        if any(x.strip() for x in current):
            parts.append(current)

    slides: list[dict[str, list[str] | str]] = []
    for part in parts or [["# Apresentação"]]:
        cleaned = [line.rstrip() for line in part if line.strip()]
        if not cleaned:
            continue
        title = _clean_markdown_text(cleaned[0]) or "Slide"
        bullets = [_clean_markdown_text(line) for line in cleaned[1:]]
        bullets = [b for b in bullets if b]
        slides.append({"title": title[:90], "bullets": bullets[:9]})
    return slides or [{"title": "Apresentação", "bullets": []}]


PPTX_W = 12192000
PPTX_H = 6858000
PPTX_PALETTE = ["4F6EF7", "10B981", "F59E0B", "EF4444", "8B5CF6", "06B6D4"]


def _pptx_p(
    text: str,
    *,
    size: int = 2200,
    color: str = "374151",
    bold: bool = False,
    align: str = "l",
    bullet: bool = False,
) -> str:
    safe = xml_escape(text[:260])
    bullet_xml = '<a:buChar char="•"/>' if bullet else ""
    mar = ' marL="342900" indent="-228600"' if bullet else ""
    bold_attr = ' b="1"' if bold else ""
    return (
        f'<a:p><a:pPr algn="{align}"{mar}>{bullet_xml}</a:pPr>'
        f'<a:r><a:rPr lang="pt-BR" sz="{size}"{bold_attr}>'
        f'<a:solidFill><a:srgbClr val="{color}"/></a:solidFill>'
        f'</a:rPr><a:t>{safe}</a:t></a:r></a:p>'
    )


def _pptx_fill(color: str | None) -> str:
    if not color:
        return "<a:noFill/>"
    return f'<a:solidFill><a:srgbClr val="{color}"/></a:solidFill>'


def _pptx_line(color: str | None = None, width: int = 12700) -> str:
    if not color:
        return "<a:ln><a:noFill/></a:ln>"
    return f'<a:ln w="{width}"><a:solidFill><a:srgbClr val="{color}"/></a:solidFill></a:ln>'


def _pptx_shape(
    shape_id: int,
    name: str,
    x: int,
    y: int,
    cx: int,
    cy: int,
    *,
    prst: str = "roundRect",
    fill: str | None = "FFFFFF",
    line: str | None = None,
    paragraphs: list[str] | None = None,
    font_size: int = 2200,
    color: str = "374151",
    bold: bool = False,
    align: str = "l",
    bullet: bool = False,
    anchor: str = "mid",
) -> str:
    body = ""
    if paragraphs is not None:
        ps = paragraphs or [" "]
        body = (
            f'<p:txBody><a:bodyPr wrap="square" anchor="{anchor}" lIns="152400" rIns="152400" tIns="91440" bIns="91440"/>'
            '<a:lstStyle/>'
            + "".join(
                _pptx_p(p, size=font_size, color=color, bold=bold, align=align, bullet=bullet)
                for p in ps
            )
            + '</p:txBody>'
        )
    return f'''<p:sp>
  <p:nvSpPr><p:cNvPr id="{shape_id}" name="{xml_escape(name)}"/><p:cNvSpPr txBox="1"/><p:nvPr/></p:nvSpPr>
  <p:spPr><a:xfrm><a:off x="{x}" y="{y}"/><a:ext cx="{cx}" cy="{cy}"/></a:xfrm><a:prstGeom prst="{prst}"><a:avLst/></a:prstGeom>{_pptx_fill(fill)}{_pptx_line(line)}</p:spPr>
  {body}
</p:sp>'''


def _pptx_title_bar(title: str, slide_no: int, accent: str) -> str:
    return (
        _pptx_shape(2, "Accent bar", 0, 0, 250000, PPTX_H, prst="rect", fill=accent)
        + _pptx_shape(3, "Slide title", 620000, 280000, 9350000, 760000, fill=None, paragraphs=[title], font_size=3000, color="111827", bold=True, anchor="mid")
        + _pptx_shape(4, "Slide number", 11270000, 6160000, 520000, 300000, fill=None, paragraphs=[str(slide_no)], font_size=1200, color="94A3B8", align="r")
    )


def _pptx_decor(shape_start: int, accent: str) -> str:
    return (
        _pptx_shape(shape_start, "Decor circle", 10600000, -470000, 1800000, 1800000, prst="ellipse", fill=accent, line=None)
        + _pptx_shape(shape_start + 1, "Decor circle 2", 11200000, 740000, 850000, 850000, prst="ellipse", fill="E0E7FF", line=None)
        + _pptx_shape(shape_start + 2, "Footer line", 620000, 6060000, 10200000, 25000, prst="rect", fill="E5E7EB", line=None)
    )


def _pptx_layout_kind(title: str, bullets: list[str], slide_no: int) -> str:
    low = title.lower()
    joined = " ".join(bullets).lower()
    if slide_no == 1:
        return "cover"
    if any(w in low for w in ("agenda", "sumário", "sumario", "roteiro", "pauta")):
        return "agenda"
    if any(w in low for w in ("perguntas", "obrigado", "conclusão", "conclusao", "próximos passos", "proximos passos")):
        return "closing"
    if any(w in low for w in ("processo", "etapas", "jornada", "fluxo", "roadmap", "cronograma", "plano de ação", "plano de acao")):
        return "process"
    if re.search(r"\b\d+[.,]?\d*\s*(%|x|k|mil|mi|m|dias?|anos?|r\$)", joined, re.IGNORECASE):
        return "metrics"
    if 2 <= len(bullets) <= 4:
        return "cards"
    return "columns"


def _pptx_cover(title: str, bullets: list[str]) -> str:
    subtitle = bullets[0] if bullets else "Apresentação profissional"
    return (
        _pptx_shape(2, "Cover band", 0, 0, PPTX_W, PPTX_H, prst="rect", fill="111827")
        + _pptx_shape(3, "Cover accent", 0, 0, 4300000, PPTX_H, prst="rect", fill="4F6EF7")
        + _pptx_shape(4, "Cover circle", 8900000, 610000, 2300000, 2300000, prst="ellipse", fill="10B981")
        + _pptx_shape(5, "Cover circle 2", 10100000, 2470000, 950000, 950000, prst="ellipse", fill="8B5CF6")
        + _pptx_shape(6, "Cover title", 670000, 1450000, 7700000, 1700000, fill=None, paragraphs=[title], font_size=4200, color="FFFFFF", bold=True, anchor="mid")
        + _pptx_shape(7, "Cover subtitle", 720000, 3300000, 6600000, 720000, fill=None, paragraphs=[subtitle], font_size=2100, color="DBEAFE", anchor="mid")
        + _pptx_shape(8, "Cover label", 720000, 5150000, 2550000, 430000, fill="FFFFFF", line=None, paragraphs=["Nexus Presentation"], font_size=1300, color="111827", bold=True, align="ctr")
    )


def _pptx_agenda(title: str, bullets: list[str], slide_no: int, accent: str) -> str:
    shapes = [_pptx_title_bar(title, slide_no, accent), _pptx_decor(40, accent)]
    for i, bullet in enumerate((bullets or ["Contexto", "Desenvolvimento", "Próximos passos"])[:7]):
        y = 1260000 + i * 590000
        color = PPTX_PALETTE[i % len(PPTX_PALETTE)]
        shapes.append(_pptx_shape(10 + i * 2, f"Agenda number {i+1}", 830000, y, 440000, 440000, prst="ellipse", fill=color, paragraphs=[str(i + 1)], font_size=1700, color="FFFFFF", bold=True, align="ctr"))
        shapes.append(_pptx_shape(11 + i * 2, f"Agenda item {i+1}", 1500000, y - 20000, 8700000, 470000, fill="F8FAFC", line="E5E7EB", paragraphs=[bullet], font_size=2000, color="1F2937"))
    return "".join(shapes)


def _pptx_cards(title: str, bullets: list[str], slide_no: int, accent: str) -> str:
    items = (bullets or ["Ideia principal"])[:4]
    shapes = [_pptx_title_bar(title, slide_no, accent), _pptx_decor(50, accent)]
    coords = [(900000, 1420000), (6350000, 1420000), (900000, 3650000), (6350000, 3650000)]
    for i, bullet in enumerate(items):
        x, y = coords[i]
        color = PPTX_PALETTE[i % len(PPTX_PALETTE)]
        shapes.append(_pptx_shape(10 + i * 4, f"Card {i+1}", x, y, 4850000, 1550000, fill="FFFFFF", line="E5E7EB"))
        shapes.append(_pptx_shape(11 + i * 4, f"Card accent {i+1}", x, y, 90000, 1550000, prst="rect", fill=color))
        shapes.append(_pptx_shape(12 + i * 4, f"Card icon {i+1}", x + 260000, y + 270000, 470000, 470000, prst="ellipse", fill=color, paragraphs=[str(i + 1)], font_size=1600, color="FFFFFF", bold=True, align="ctr"))
        shapes.append(_pptx_shape(13 + i * 4, f"Card text {i+1}", x + 900000, y + 230000, 3550000, 900000, fill=None, paragraphs=[bullet], font_size=1900, color="1F2937", bold=True, anchor="mid"))
    return "".join(shapes)


def _pptx_columns(title: str, bullets: list[str], slide_no: int, accent: str) -> str:
    items = bullets or ["Conteúdo principal"]
    mid = (len(items) + 1) // 2
    shapes = [_pptx_title_bar(title, slide_no, accent), _pptx_decor(60, accent)]
    for col, col_items in enumerate((items[:mid], items[mid:])):
        x = 870000 + col * 5400000
        shapes.append(_pptx_shape(10 + col, f"Column bg {col}", x, 1300000, 4850000, 4450000, fill="F8FAFC", line="E5E7EB"))
        for i, bullet in enumerate(col_items[:7]):
            y = 1540000 + i * 560000
            color = PPTX_PALETTE[(i + col * 3) % len(PPTX_PALETTE)]
            shapes.append(_pptx_shape(30 + col * 20 + i * 2, f"Bullet dot {col}-{i}", x + 330000, y + 95000, 170000, 170000, prst="ellipse", fill=color))
            shapes.append(_pptx_shape(31 + col * 20 + i * 2, f"Bullet text {col}-{i}", x + 630000, y - 20000, 3850000, 360000, fill=None, paragraphs=[bullet], font_size=1750, color="374151"))
    return "".join(shapes)


def _metric_parts(text: str) -> tuple[str, str]:
    match = re.search(r"([+-]?\d+[.,]?\d*\s*(?:%|x|k|mil|mi|m|dias?|anos?|r\$)?)", text, re.IGNORECASE)
    if not match:
        return "•", text
    value = match.group(1).strip()
    label = (text[:match.start()] + text[match.end():]).strip(" -:–")
    return value, label or text


def _pptx_metrics(title: str, bullets: list[str], slide_no: int, accent: str) -> str:
    items = (bullets or ["100% Resultado esperado"])[:4]
    shapes = [_pptx_title_bar(title, slide_no, accent), _pptx_decor(70, accent)]
    card_w = 2500000
    for i, bullet in enumerate(items):
        value, label = _metric_parts(bullet)
        x = 850000 + i * 2780000
        color = PPTX_PALETTE[i % len(PPTX_PALETTE)]
        shapes.append(_pptx_shape(10 + i * 4, f"Metric card {i}", x, 1900000, card_w, 2550000, fill="FFFFFF", line="E5E7EB"))
        shapes.append(_pptx_shape(11 + i * 4, f"Metric top {i}", x, 1900000, card_w, 160000, prst="rect", fill=color))
        shapes.append(_pptx_shape(12 + i * 4, f"Metric value {i}", x + 180000, 2390000, card_w - 360000, 850000, fill=None, paragraphs=[value], font_size=3600, color=color, bold=True, align="ctr"))
        shapes.append(_pptx_shape(13 + i * 4, f"Metric label {i}", x + 280000, 3420000, card_w - 560000, 700000, fill=None, paragraphs=[label], font_size=1600, color="374151", align="ctr", anchor="mid"))
    return "".join(shapes)


def _pptx_process(title: str, bullets: list[str], slide_no: int, accent: str) -> str:
    items = (bullets or ["Etapa 1", "Etapa 2", "Etapa 3"])[:5]
    shapes = [_pptx_title_bar(title, slide_no, accent), _pptx_decor(80, accent)]
    step_w = 2050000
    for i, bullet in enumerate(items):
        x = 700000 + i * 2200000
        color = PPTX_PALETTE[i % len(PPTX_PALETTE)]
        prst = "chevron" if i < len(items) - 1 else "roundRect"
        shapes.append(_pptx_shape(10 + i * 3, f"Process step {i}", x, 2500000, step_w, 1350000, prst=prst, fill=color, line=None, paragraphs=[bullet], font_size=1550, color="FFFFFF", bold=True, align="ctr", anchor="mid"))
        shapes.append(_pptx_shape(11 + i * 3, f"Process number {i}", x + 120000, 2160000, 440000, 440000, prst="ellipse", fill="FFFFFF", line=color, paragraphs=[str(i + 1)], font_size=1500, color=color, bold=True, align="ctr"))
    return "".join(shapes)


def _pptx_closing(title: str, bullets: list[str], slide_no: int, accent: str) -> str:
    lines = bullets[:3] or ["Obrigado"]
    return (
        _pptx_shape(2, "Closing bg", 0, 0, PPTX_W, PPTX_H, prst="rect", fill="F8FAFC")
        + _pptx_shape(3, "Closing accent", 0, 0, PPTX_W, 900000, prst="rect", fill=accent)
        + _pptx_shape(4, "Closing circle", 8800000, 3500000, 2100000, 2100000, prst="ellipse", fill="E0E7FF")
        + _pptx_shape(5, "Closing title", 1050000, 2050000, 7900000, 1150000, fill=None, paragraphs=[title], font_size=4000, color="111827", bold=True, align="ctr", anchor="mid")
        + _pptx_shape(6, "Closing bullets", 2300000, 3450000, 5600000, 1250000, fill="FFFFFF", line="E5E7EB", paragraphs=lines, font_size=1850, color="374151", align="ctr", anchor="mid")
        + _pptx_shape(7, "Slide number", 11270000, 6160000, 520000, 300000, fill=None, paragraphs=[str(slide_no)], font_size=1200, color="94A3B8", align="r")
    )


def _pptx_slide_xml(title: str, bullets: list[str], slide_no: int) -> str:
    accent = PPTX_PALETTE[(slide_no - 1) % len(PPTX_PALETTE)]
    kind = _pptx_layout_kind(title, bullets, slide_no)
    if kind == "cover":
        content = _pptx_cover(title, bullets)
    elif kind == "agenda":
        content = _pptx_agenda(title, bullets, slide_no, accent)
    elif kind == "process":
        content = _pptx_process(title, bullets, slide_no, accent)
    elif kind == "metrics":
        content = _pptx_metrics(title, bullets, slide_no, accent)
    elif kind == "cards":
        content = _pptx_cards(title, bullets, slide_no, accent)
    elif kind == "closing":
        content = _pptx_closing(title, bullets, slide_no, accent)
    else:
        content = _pptx_columns(title, bullets, slide_no, accent)
    return f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<p:sld xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" xmlns:p="http://schemas.openxmlformats.org/presentationml/2006/main">
  <p:cSld>
    <p:bg><p:bgPr><a:solidFill><a:srgbClr val="F8FAFC"/></a:solidFill><a:effectLst/></p:bgPr></p:bg>
    <p:spTree>
      <p:nvGrpSpPr><p:cNvPr id="1" name=""/><p:cNvGrpSpPr/><p:nvPr/></p:nvGrpSpPr>
      <p:grpSpPr><a:xfrm><a:off x="0" y="0"/><a:ext cx="0" cy="0"/><a:chOff x="0" y="0"/><a:chExt cx="0" cy="0"/></a:xfrm></p:grpSpPr>
      {content}
    </p:spTree>
  </p:cSld>
  <p:clrMapOvr><a:masterClrMapping/></p:clrMapOvr>
</p:sld>'''


def _pptx_content_types(slide_count: int) -> str:
    slide_overrides = "".join(
        f'<Override PartName="/ppt/slides/slide{i}.xml" ContentType="application/vnd.openxmlformats-officedocument.presentationml.slide+xml"/>'
        for i in range(1, slide_count + 1)
    )
    return f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
  <Default Extension="xml" ContentType="application/xml"/>
  <Override PartName="/ppt/presentation.xml" ContentType="application/vnd.openxmlformats-officedocument.presentationml.presentation.main+xml"/>
  <Override PartName="/ppt/slideMasters/slideMaster1.xml" ContentType="application/vnd.openxmlformats-officedocument.presentationml.slideMaster+xml"/>
  <Override PartName="/ppt/slideLayouts/slideLayout1.xml" ContentType="application/vnd.openxmlformats-officedocument.presentationml.slideLayout+xml"/>
  <Override PartName="/ppt/theme/theme1.xml" ContentType="application/vnd.openxmlformats-officedocument.theme+xml"/>
  <Override PartName="/docProps/core.xml" ContentType="application/vnd.openxmlformats-package.core-properties+xml"/>
  <Override PartName="/docProps/app.xml" ContentType="application/vnd.openxmlformats-officedocument.extended-properties+xml"/>
  {slide_overrides}
</Types>'''


def _create_pptx(content: str, path: Path) -> None:
    slides = _split_pptx_slides(content)
    slide_count = len(slides)
    sld_ids = "".join(
        f'<p:sldId id="{255 + i}" r:id="rId{i + 1}"/>'
        for i in range(1, slide_count + 1)
    )
    pres_rels = (
        '<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/slideMaster" Target="slideMasters/slideMaster1.xml"/>'
        + "".join(
            f'<Relationship Id="rId{i + 1}" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/slide" Target="slides/slide{i}.xml"/>'
            for i in range(1, slide_count + 1)
        )
    )

    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("[Content_Types].xml", _pptx_content_types(slide_count))
        zf.writestr("_rels/.rels", '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="ppt/presentation.xml"/>
  <Relationship Id="rId2" Type="http://schemas.openxmlformats.org/package/2006/relationships/metadata/core-properties" Target="docProps/core.xml"/>
  <Relationship Id="rId3" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/extended-properties" Target="docProps/app.xml"/>
</Relationships>''')
        zf.writestr("docProps/core.xml", '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<cp:coreProperties xmlns:cp="http://schemas.openxmlformats.org/package/2006/metadata/core-properties" xmlns:dc="http://purl.org/dc/elements/1.1/"><dc:title>Apresentação</dc:title><dc:creator>Nexus</dc:creator></cp:coreProperties>''')
        zf.writestr("docProps/app.xml", f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Properties xmlns="http://schemas.openxmlformats.org/officeDocument/2006/extended-properties"><Application>Nexus</Application><Slides>{slide_count}</Slides></Properties>''')
        zf.writestr("ppt/presentation.xml", f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<p:presentation xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" xmlns:p="http://schemas.openxmlformats.org/presentationml/2006/main">
  <p:sldMasterIdLst><p:sldMasterId id="2147483648" r:id="rId1"/></p:sldMasterIdLst>
  <p:sldIdLst>{sld_ids}</p:sldIdLst>
  <p:sldSz cx="12192000" cy="6858000" type="wide"/>
  <p:notesSz cx="6858000" cy="9144000"/>
</p:presentation>''')
        zf.writestr("ppt/_rels/presentation.xml.rels", f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">{pres_rels}</Relationships>''')
        zf.writestr("ppt/slideMasters/slideMaster1.xml", '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<p:sldMaster xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" xmlns:p="http://schemas.openxmlformats.org/presentationml/2006/main"><p:cSld><p:spTree><p:nvGrpSpPr><p:cNvPr id="1" name=""/><p:cNvGrpSpPr/><p:nvPr/></p:nvGrpSpPr><p:grpSpPr/></p:spTree></p:cSld><p:clrMap bg1="lt1" tx1="dk1" bg2="lt2" tx2="dk2" accent1="accent1" accent2="accent2" accent3="accent3" accent4="accent4" accent5="accent5" accent6="accent6" hlink="hlink" folHlink="folHlink"/><p:sldLayoutIdLst><p:sldLayoutId id="2147483649" r:id="rId1"/></p:sldLayoutIdLst><p:txStyles><p:titleStyle/><p:bodyStyle/><p:otherStyle/></p:txStyles></p:sldMaster>''')
        zf.writestr("ppt/slideMasters/_rels/slideMaster1.xml.rels", '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"><Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/slideLayout" Target="../slideLayouts/slideLayout1.xml"/><Relationship Id="rId2" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/theme" Target="../theme/theme1.xml"/></Relationships>''')
        zf.writestr("ppt/slideLayouts/slideLayout1.xml", '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<p:sldLayout xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" xmlns:p="http://schemas.openxmlformats.org/presentationml/2006/main" type="blank" preserve="1"><p:cSld name="Blank"><p:spTree><p:nvGrpSpPr><p:cNvPr id="1" name=""/><p:cNvGrpSpPr/><p:nvPr/></p:nvGrpSpPr><p:grpSpPr/></p:spTree></p:cSld><p:clrMapOvr><a:masterClrMapping/></p:clrMapOvr></p:sldLayout>''')
        zf.writestr("ppt/slideLayouts/_rels/slideLayout1.xml.rels", '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"><Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/slideMaster" Target="../slideMasters/slideMaster1.xml"/></Relationships>''')
        zf.writestr("ppt/theme/theme1.xml", '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<a:theme xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" name="Nexus"><a:themeElements><a:clrScheme name="Nexus"><a:dk1><a:srgbClr val="111827"/></a:dk1><a:lt1><a:srgbClr val="FFFFFF"/></a:lt1><a:dk2><a:srgbClr val="374151"/></a:dk2><a:lt2><a:srgbClr val="F3F4F6"/></a:lt2><a:accent1><a:srgbClr val="4F6EF7"/></a:accent1><a:accent2><a:srgbClr val="10B981"/></a:accent2><a:accent3><a:srgbClr val="F59E0B"/></a:accent3><a:accent4><a:srgbClr val="EF4444"/></a:accent4><a:accent5><a:srgbClr val="8B5CF6"/></a:accent5><a:accent6><a:srgbClr val="06B6D4"/></a:accent6><a:hlink><a:srgbClr val="2563EB"/></a:hlink><a:folHlink><a:srgbClr val="7C3AED"/></a:folHlink></a:clrScheme><a:fontScheme name="Nexus"><a:majorFont><a:latin typeface="Aptos Display"/></a:majorFont><a:minorFont><a:latin typeface="Aptos"/></a:minorFont></a:fontScheme><a:fmtScheme name="Nexus"><a:fillStyleLst><a:solidFill><a:schemeClr val="phClr"/></a:solidFill></a:fillStyleLst><a:lnStyleLst><a:ln w="6350"><a:solidFill><a:schemeClr val="phClr"/></a:solidFill></a:ln></a:lnStyleLst><a:effectStyleLst><a:effectStyle><a:effectLst/></a:effectStyle></a:effectStyleLst><a:bgFillStyleLst><a:solidFill><a:schemeClr val="phClr"/></a:solidFill></a:bgFillStyleLst></a:fmtScheme></a:themeElements></a:theme>''')
        for i, slide in enumerate(slides, start=1):
            zf.writestr(
                f"ppt/slides/slide{i}.xml",
                _pptx_slide_xml(str(slide["title"]), list(slide["bullets"]), i),
            )
            zf.writestr(f"ppt/slides/_rels/slide{i}.xml.rels", '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"><Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/slideLayout" Target="../slideLayouts/slideLayout1.xml"/></Relationships>''')


def _mime(ext: str) -> str:
    return {
        ".pdf":  "application/pdf",
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        ".pptx": "application/vnd.openxmlformats-officedocument.presentationml.presentation",
        ".zip":  "application/zip",
        ".py":   "text/x-python",
        ".js":   "text/javascript",
        ".ts":   "text/typescript",
        ".html": "text/html",
        ".css":  "text/css",
        ".json": "application/json",
        ".md":   "text/markdown",
        ".csv":  "text/csv",
        ".xml":  "application/xml",
        ".yaml": "text/yaml",
        ".yml":  "text/yaml",
        ".sh":   "text/x-sh",
        ".sql":  "text/x-sql",
    }.get(ext, "text/plain")


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def generate_file(filename: str, content: str, workspace_dir: Path) -> dict[str, Any]:
    """Create a file on disk from AI-provided content. Returns metadata."""
    workspace_dir.mkdir(parents=True, exist_ok=True)
    ext = Path(filename).suffix.lower()
    stored = _safe_name(filename)

    # Normalize .xls → .xlsx
    if ext == ".xls":
        stored = Path(stored).stem + ".xlsx"
        ext = ".xlsx"

    dest = workspace_dir / stored

    try:
        if ext == ".docx":
            _create_docx(content, dest)
        elif ext == ".xlsx":
            _create_xlsx(content, dest)
        elif ext == ".pptx":
            _create_pptx(content, dest)
        elif ext == ".pdf":
            _create_pdf(content, dest)
        else:
            _create_text(content, dest)

        return {
            "filename": filename, "stored_name": stored,
            "size": dest.stat().st_size, "mime": _mime(ext), "ok": True,
        }
    except Exception as exc:
        return {"filename": filename, "ok": False, "error": str(exc)}


def bundle_zip(pairs: list[tuple[str, Path]], zip_path: Path) -> None:
    """Bundle list of (arcname, path) into a ZIP file."""
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for arcname, src in pairs:
            if src.exists():
                zf.write(src, arcname)
